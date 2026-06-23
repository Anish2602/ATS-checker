import re
import docx
from typing import Dict, Any, List

class ResumeParserService:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extracts text stream from PDF bytes using pypdf."""
        import io
        from pypdf import PdfReader
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            return text
        except Exception:
            return ""

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extracts text from DOCX bytes using python-docx."""
        import io
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)

    @staticmethod
    def clean_text(text: str) -> str:
        """Cleans control characters, ligatures, and normalizes spacing."""
        text = text.replace("\u2013", "-").replace("\u2014", "-")
        text = text.replace("\ufb01", "fi").replace("\ufb02", "fl")
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    @classmethod
    def parse_resume(cls, file_bytes: bytes, file_type: str) -> Dict[str, Any]:
        """Parses a resume, returning clean text and section segmentation."""
        raw_text = ""
        if file_type == "pdf":
            raw_text = cls.extract_text_from_pdf(file_bytes)
        elif file_type in ["docx", "doc"]:
            raw_text = cls.extract_text_from_docx(file_bytes)
        else:
            raw_text = file_bytes.decode("utf-8", errors="ignore")

        cleaned_text = cls.clean_text(raw_text)
        
        # Section classification boundaries
        sections = cls.segment_sections(cleaned_text)
        
        # Detect formatting elements
        formatting_flags = cls.detect_formatting_issues(file_bytes, file_type)
        
        return {
            "raw_text": cleaned_text,
            "sections": sections,
            "formatting": formatting_flags
        }

    @staticmethod
    def segment_sections(text: str) -> Dict[str, str]:
        """Segments text into main resume sections based on typical keywords."""
        sections = {
            "summary": "",
            "experience": "",
            "skills": "",
            "projects": "",
            "education": ""
        }
        
        # Standard headings patterns
        heading_patterns = {
            "summary": r'\b(professional summary|summary|about me|profile|objective)\b',
            "experience": r'\b(experience|work history|employment|professional background|experience history)\b',
            "skills": r'\b(skills|technical skills|technologies|core competencies|expertise)\b',
            "projects": r'\b(projects|academic projects|key projects|personal projects)\b',
            "education": r'\b(education|academic credentials|qualifications|academic history)\b'
        }
        
        # Build index positions of headers
        matches = []
        for section, pattern in heading_patterns.items():
            for match in re.finditer(pattern, text, re.IGNORECASE):
                matches.append((match.start(), match.end(), section))
        
        # Sort matches by start index
        matches.sort(key=lambda x: x[0])
        
        # If no sections are detected, dump all text in experience as a fallback
        if not matches:
            sections["experience"] = text
            return sections
        
        # Extract text blocks between headers
        for i in range(len(matches)):
            start_pos = matches[i][1]
            end_pos = matches[i+1][0] if i + 1 < len(matches) else len(text)
            section_name = matches[i][2]
            sections[section_name] += " " + text[start_pos:end_pos].strip()
            
        # Cleanup
        for k in sections:
            sections[k] = sections[k].strip()
            
        return sections

    @staticmethod
    def detect_formatting_issues(file_bytes: bytes, file_type: str) -> Dict[str, Any]:
        """Runs simple checks to detect tables, complex shapes, or column layouts."""
        issues = {
            "has_tables": False,
            "has_columns": False,
            "has_images": False,
            "is_scanned": False,
            "page_count": 1
        }
        
        if file_type == "pdf":
            try:
                import io
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                issues["page_count"] = len(reader.pages)
                
                total_text_length = 0
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        total_text_length += len(t)
                        # Heuristic column detection in text lines: double spaces or tab characters
                        if "   " in t or "\t" in t:
                            issues["has_columns"] = True
                    
                    # Check for images/objects in pypdf
                    if page.images and len(page.images) > 0:
                        issues["has_images"] = True
                        
                # Scanned check: very little text relative to page count
                if len(reader.pages) > 0 and total_text_length / len(reader.pages) < 100:
                    issues["is_scanned"] = True
                    
            except Exception:
                pass
                
        elif file_type == "docx":
            try:
                import io
                doc = docx.Document(io.BytesIO(file_bytes))
                if len(doc.tables) > 0:
                    issues["has_tables"] = True
            except Exception:
                pass
                
        return issues
